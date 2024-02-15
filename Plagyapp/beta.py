from __future__ import print_function
import sqlite3
from docx import Document
import tkinter as tk
import bcrypt
from tkinter import messagebox
import tkinter.messagebox
import requests
import json
from PIL import Image, ImageTk
import webbrowser
import ProWritingAidSDK
from ProWritingAidSDK.rest import ApiException
from tkinter import filedialog, messagebox

plag_per = ""
source_urls = []
status = 0

conn = sqlite3.connect('data.db')
cursor = conn.cursor()

cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
    username TEXT NOT NULL,
    password TEXT NOT NULL)''')


def signup():
    username = username_entry.get()
    password = password_entry.get()

    if username != '' and password != '':
        cursor.execute("SELECT username FROM users WHERE username = ?", [username])
        if cursor.fetchone() is not None:
            messagebox.showerror("ERROR", 'Username already exists.')
        else:
            encoded_password = password.encode('utf-8')
            hashed_password = bcrypt.hashpw(encoded_password, bcrypt.gensalt())
            print(hashed_password)
            cursor.execute("INSERT INTO users VALUES (?, ?)", [username, hashed_password])
            conn.commit()
            messagebox.showinfo("Success", 'Account has been created.')
    else:
        messagebox.showerror("Error", "Enter all data.")


def login_account():
    username = username_entry2.get()
    password = password_entry2.get()

    if username != '' and password != '':
        cursor.execute('SELECT password FROM users WHERE username=?', [username])
        result = cursor.fetchone()
        if result:
            if bcrypt.checkpw(password.encode('utf8'), result[0]):
                messagebox.showinfo("Success", "Logged in successfully")
                hide_login()
                show_main_app()
            else:
                messagebox.showerror("ERROR", "Invalid pwd")
        else:
            messagebox.showerror("Error", "Invalid Username")
    else:
        messagebox.showerror("Error", "Enter all data")


def login():
    hide_signup()
    global username_entry2, password_entry2, login_label2, username_label2, password_label2, login_btn2
    show_login()


def show_signup():
    Signup_label.place(x=325, y=190)
    username_label.place(x=170, y=290)
    password_label.place(x=170, y=380)
    username_entry.place(x=340, y=290)
    password_entry.place(x=340, y=380)
    login_label.place(x=220, y=550)
    login_btn.place(x=480, y=545)
    signup_btn.place(x=325, y=450)


def hide_signup():
    Signup_label.place_forget()
    username_label.place_forget()
    password_label.place_forget()
    username_entry.place_forget()
    password_entry.place_forget()
    login_label.place_forget()
    login_btn.place_forget()
    signup_btn.place_forget()


def show_login():
    login_label2.place(x=325, y=190)
    username_label2.place(x=170, y=290)
    password_label2.place(x=170, y=380)
    username_entry2.place(x=340, y=290)
    password_entry2.place(x=340, y=380)
    login_btn2.place(x=325, y=450)


def hide_login():
    login_label2.place_forget()
    username_label2.place_forget()
    password_label2.place_forget()
    username_entry2.place_forget()
    password_entry2.place_forget()
    login_btn2.place_forget()


def show_main_app():
    # Your existing UI code here
    Heading.place(x=250, y=110)
    text_box.place(x=67, y=170)
    chk_btn.place(x=330, y=400)
    reset_btn.place(x=500, y=434)
    download_rep_btn.place(x=500, y=400)
    Grammar_fix_btn.place(x=190, y=415)
    plag_percent.place(x=90, y=500)
    Links_label.place(x=90, y=550)
    Links_res.place(x=90, y=580)


def update_plag_percent_label():
    plag_percent.config(text=f"PLAGIARISM %: {plag_per}%")


def validate_textbox():
    global data
    global plag_per, source_urls, status

    content = text_box.get("1.0", "end-1c").strip()  # strip trims whitespace
    if not content:
        tkinter.messagebox.showerror(title="ERROR", message="Please enter some text before checking for plagiarism.")
        return
    word_count = len(content.split())

    if word_count > 699:
        tkinter.messagebox.showerror(title="ERROR", message="700 words limit exceeds")
    else:
        url = "https://plagiarism-checker-and-auto-citation-generator-multi-lingual.p.rapidapi.com/plagiarism"

        payload = {
            "text": content,
            "language": "en",
            "includeCitations": False,
            "scrapeSources": False
        }
        headers = {
            "content-type": "application/json",
            "X-RapidAPI-Key": "5c32135ac8mshcdd8ec5e3e2d725p13bc44jsne6eeb109d50f",
            "X-RapidAPI-Host": "plagiarism-checker-and-auto-citation-generator-multi-lingual.p.rapidapi.com"
        }

        response = requests.post(url, json=payload, headers=headers)
        status = response.status_code
        print(f"status : {status}")  # TESTING ONLY

        json_data = response.content
        data = json.loads(json_data)
        print(data)

        plag_per = data["percentPlagiarism"]
        source_urls = [source["url"] for source in data["sources"]]

        update_plag_percent_label()
        update_links()
        enable_generate_word_button()


def generate_word_report():
    try:
        word_filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if word_filename:
            document = Document()
            document.add_heading("PLAGIARISM REPORT", level=1)
            document.add_paragraph(f"Plagiarism Percentage: {plag_per}%")
            document.add_heading("Source URLs:", level=2)

            for i, (source_url, matches) in enumerate(zip(source_urls, data["sources"]), start=1):
                document.add_paragraph(f"{i}. {source_url}", style='ListBullet')

                for j, match in enumerate(matches["matches"], start=1):
                    match_text = match.get("matchText", "")
                    document.add_paragraph(f"  - {match_text}")

            document.save(word_filename)
            messagebox.showinfo("Word Report", f"Word Report generated successfully and saved as: {word_filename}")
    except Exception as e:
        messagebox.showerror("Error", f"Error generating Word Report")


def update_links():
    Links_res.config(state=tk.NORMAL)
    Links_res.delete("1.0", tk.END)  # Clear previous content

    for i, source_url in enumerate(source_urls, start=1):
        link_text = f"{i}. {source_url}\n"
        Links_res.insert(tk.END, link_text, "hyperlink")
        Links_res.tag_add("hyperlink", f"{i}.0", f"{i}.{len(source_url) + 2}")
        Links_res.tag_config("hyperlink", foreground="blue", underline=True)
        Links_res.tag_bind("hyperlink", "<Button-1>", lambda event, url=source_url: webbrowser.open(url))

    Links_res.config(state=tk.DISABLED)


def enable_generate_word_button():
    download_rep_btn.config(state=tk.NORMAL, bg="#C1FF72")


def disable_generate_word_button():
    download_rep_btn.config(state=tk.DISABLED, bg="Grey")


def fix_grammar():
    global plag_per, source_urls, status

    content = text_box.get("1.0", "end-1c")
    if not content:
        tkinter.messagebox.showerror(title="ERROR", message="Please enter some text before fixing grammar mistakes.")
        return

    configuration = ProWritingAidSDK.Configuration()
    configuration.host = 'https://api.prowritingaid.com'
    configuration.api_key['licenseCode'] = '1D459A39-464C-4176-BFA5-D1C24FF44915'

    api_instance = ProWritingAidSDK.TextApi(ProWritingAidSDK.ApiClient('https://api.prowritingaid.com'))

    wrong_sent = content

    try:
        api_request = ProWritingAidSDK.TextAnalysisRequest(wrong_sent,
                                                           ["grammar"],
                                                           "General",
                                                           "en")
        api_response = api_instance.post(api_request)
        print(api_response)

    except ApiException as e:
        print("Exception when calling TextAnalysisRequest->get: %s\n" % e)

    tags = api_response.result.tags
    correct_sentence_chars = list(wrong_sent)

    for tag in reversed(tags):
        if tag.suggestions:
            replacement = '' if tag.suggestions[0] == '(omit)' else tag.suggestions[0]
            correct_sentence_chars[tag.start_pos:tag.end_pos + 1] = replacement

    correct_sentence = ''.join(correct_sentence_chars)

    # Update the content in the text_box
    text_box.delete("1.0", tk.END)
    text_box.insert(tk.END, correct_sentence)


def reset_application():
    global plag_per, source_urls, status
    plag_per = ""
    source_urls = []
    status = 0

    text_box.delete("1.0", tk.END)

    update_plag_percent_label()
    update_links()
    disable_generate_word_button()

    # ------------------------ UI -----------------------------------------#


# ---------------------------- WINDOW --------------------------

window = tk.Tk()
window.geometry("775x700")
window.title("PLAGIARISM CHECKER ©️")
window.config(bg="BLACK")

# Signup UI elements
Signup_label = tk.Label(window, text="Signup", bg='black', fg='#C1FF72', font="Helvetica 25 bold")
username_label = tk.Label(window, text="Username:", bg='black', fg='#C1FF72', font=("Helvetica", 15, "bold"))
password_label = tk.Label(window, text="Password:", bg='black', fg='#C1FF72', font=("Helvetica", 15, "bold"))
username_entry = tk.Entry(window, width=26, font=("Helvetica", 15, "bold"), bg="grey")
password_entry = tk.Entry(window, width=26, font=("Helvetica", 15, "bold"), show="*", bg="grey")
login_label = tk.Label(window, text="Already have an account?", bg='black', fg='#C1FF72',
                       font=("Helvetica", 15, "bold"))
login_btn = tk.Button(window, text="Login", height=2, width=14, font=("Helvetica", 10, "bold"),
                      relief=tk.RAISED, bg="black", fg="#C1FF72", command=login)
signup_btn = tk.Button(window, text="Signup", height=2, width=14, font=("Helvetica", 10, "bold"),
                       relief=tk.RAISED, bg="#C1FF72", command=signup)

# LOGIN UI ELEMENTS

login_label2 = tk.Label(window, text="Login", bg='black', fg='#C1FF72', font="Helvetica 25 bold")
username_label2 = tk.Label(window, text="Username:", bg='black', fg='#C1FF72', font=("Helvetica", 15, "bold"))
password_label2 = tk.Label(window, text="Password:", bg='black', fg='#C1FF72', font=("Helvetica", 15, "bold"))
username_entry2 = tk.Entry(window, width=26, font=("Helvetica", 15, "bold"), bg="grey")
password_entry2 = tk.Entry(window, width=26, font=("Helvetica", 15, "bold"), show="*", bg="grey")
login_btn2 = tk.Button(window, text="Login", height=2, width=14, font=("Helvetica", 10, "bold"),
                       relief=tk.RAISED, bg="#C1FF72", command=login_account)

# ------------------------- NAVBAR ---------------------------------

Photo_Open = Image.open("navbar2.png")
photo_final = ImageTk.PhotoImage(Photo_Open)
img_label = tk.Label(window, image=photo_final)
img_label.place(x=-2)

# ------------------------ UI-------------------------------=
Heading = tk.Label(window, text="CHECK FOR PLAGIARISM", bg='black', fg='#C1FF72', font="Helvetica 18 bold")
text_box = tk.Text(window, height=12, width=80, bg="#A9A9A9", wrap="word")
chk_btn = tk.Button(window, text="CHECK", height=3, width=14, font=("Helvetica", 10, "bold"), relief=tk.RAISED,
                    bg="#C1FF72", command=validate_textbox)
reset_btn = tk.Button(window, text="Reset", height=2, width=17, font=("Helvetica", 7, "bold"), relief=tk.RAISED,
                      bg="#C1FF72", command=reset_application)
download_rep_btn = tk.Button(window, text="Download Report", height=2, width=17, font=("Helvetica", 7, "bold"),
                             relief=tk.RAISED, bg="grey", command=generate_word_report, state=tk.DISABLED)
Grammar_fix_btn = tk.Button(window, text="Fix Grammar mistakes", height=2, width=17, font=("Helvetica", 7, "bold"),
                            bg="#C1FF72", command=fix_grammar)
plag_percent = tk.Label(window, text=f"PLAGIARISM % : {plag_per}", bg='black', fg='#C1FF72',
                        font=("Helvetica", 10, "bold"))
Links_label = tk.Label(window, text="LINKS : ", bg='black', fg='#C1FF72', font=("Helvetica", 10, "bold"))
Links_res = tk.Text(window, height=5, width=100, bg='black', fg='#C1FF72', font=("Helvetica", 8, "bold"))

show_signup()
window.resizable(False, False)
window.mainloop()
